import os

import docx
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.article_form.article_form_request import ArticleFormDomain
from app.models.file.article_form_request import FileDomain


class DocService:

    def __init__(self):
        self.document = docx.Document()

    def create_formatted_docx(
        self,
        article: ArticleFormDomain,
        attached_text: str,
        filepath: str,
        filename: str,
    ) -> FileDomain:
        self.document = docx.Document()

        # настройки страницы
        self._set_page_settings()

        # Задаем параметры абзаца
        self._set_paragraph_settings()

        # создание абзацев
        self._create_paragraph()

        # вставка УДК слева
        self._set_udk(article)

        # вставка блока с авторами
        self._set_authors(article)

        # вставка названия статьи
        self._add_title(article)

        # вставка аннотации
        self._add_anotation(article)

        # вставка ключевых слов
        self._add_keywords(article)

        # вставка текста
        self._add_text_body(attached_text)

        # создание списка литературы
        self._add_list_of_references(article)

        # сохранение документа
        return self._save_doc(
            filepath=filepath,
            filename=filename,
        )

    def _set_page_settings(self,):
        section = self.document.sections[0]
        section.page_height = docx.shared.Cm(29.7)
        section.page_width = docx.shared.Cm(21)
        section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
        section.top_margin = docx.shared.Cm(1.4)
        section.bottom_margin = docx.shared.Cm(2)
        section.left_margin = docx.shared.Cm(1.7)
        section.right_margin = docx.shared.Cm(1.7)

    def _set_paragraph_settings(self,):
        style = self.document.styles['Normal']
        style.paragraph_format.space_before = docx.shared.Pt(8.4)
        style.paragraph_format.space_after = docx.shared.Pt(0)
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.first_line_indent = docx.shared.Cm(0.7)
        style.font.name = 'Times New Roman'
        style.font.size = docx.shared.Pt(10)
        style.font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    def _create_paragraph(self,):
        paragraph = self.document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.first_line_indent = Pt(7)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # настройки междустрочного интервала и выравнивания текста
        paragraph_format.line_spacing_rule = 0
        paragraph_format.line_spacing = 1
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _set_udk(self, article):
        udk_block = self.document.add_paragraph()
        udk_block_alignment = udk_block.paragraph_format
        udk_block_alignment.alignment = WD_ALIGN_PARAGRAPH.LEFT
        udk_block.add_run(article.udc)

    def _set_authors(self, article: ArticleFormDomain):
        # TODO:
        authors_block = self.document.add_paragraph()
        authors_block_format = authors_block.paragraph_format
        authors_block_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        authors_names = []
        for author in article.authors:
            author_info = author.place_of_study if author.place_of_study else author.place_of_work
            name = f'{author.last_name} {author.first_name[0].upper()}. {author.surname[0].upper()}.'
            authors_names.append(f'{name}, студент бакалавр\n{author_info}\n')

        scientific_adviser_name = self._format_scientific_adviser_name(
            article.scientific_adviser_fullname
        )

        authors_workplace = f'{article.scientific_adviser_academic_degree}, ' \
                            f'{scientific_adviser_name}\n' \
                            f'{article.scientific_adviser_institute}'

        authors_block.add_run(''.join(authors_names)).bold = True
        authors_block.add_run('Научный руководитель:\n').bold = True
        authors_block.add_run(authors_workplace).bold = True

    def _add_title(self, article: ArticleFormDomain):
        title_block = self.document.add_paragraph()
        title_block_format = title_block.paragraph_format
        title_block_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        title_block.add_run(article.title_rus.upper()).bold = True
        title_block.add_run('\n' + article.title_eng.upper()).bold = True

    def _add_anotation(self, article: ArticleFormDomain):
        abstract_block = self.document.add_paragraph()
        abstract_format = abstract_block.paragraph_format
        abstract_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        abstract_format.left_indent = Cm(0.7)

        abstract_name_rus = 'Аннотация'
        abstract_name_eng = '\tAnnotation'

        abstract_font_rus = abstract_block.add_run(abstract_name_rus + '\n').font
        text_1 = abstract_block.add_run(article.abstract_rus + '\n').font
        abstract_font_eng = abstract_block.add_run(abstract_name_eng + '\n').font
        text_2 = abstract_block.add_run(article.abstract_eng).font

        abstract_font_rus.italic = True
        abstract_font_eng.italic = True
        text_1.italic = True
        text_2.italic = True

        abstract_font_rus.size = Pt(9)
        abstract_font_eng.size = Pt(9)
        text_1.size = Pt(9)
        text_2.size = Pt(9)

    def _add_keywords(self, article: ArticleFormDomain):
        keywords_block = self.document.add_paragraph()
        keywords_format = keywords_block.paragraph_format
        keywords_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        keywords_format.left_indent = Cm(0.7)

        keywords_font = keywords_block.add_run(
            f'Ключевые слова:\n{article.keywords_rus}\n'
            f'\tKeywords:\n{article.keywords_eng}'
        ).font

        keywords_font.italic = True
        keywords_font.size = Pt(9)

    def _add_text_body(self, attached_text: str):
        text = self.document.add_paragraph(style='Normal')
        text.add_run(attached_text)

    def _add_list_of_references(self, article: ArticleFormDomain):
        list_of_references = self.document.add_paragraph()
        list_of_references.add_run(
            f"Список литературы\n{article.list_of_references}"
        ).bold = True

    def _save_doc(
        self,
        filepath: str,
        filename: str,
    ) -> FileDomain:
        path = os.path.join(filepath, filename)
        self.document.save(path)
        return FileDomain(
            path=path,
            name=filename,
            size_kb=int(os.stat(path).st_size / 1024),
        )

    def _format_scientific_adviser_name(self, name: str) -> str:
        name_list = name.split(' ')
        if not len(name_list) == 3:
            return name
        return f'{name_list[0]} {name_list[1][0].upper()}. {name_list[2][0].upper()}.'
